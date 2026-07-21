"""Build the SafeAgent project guide DOCX with a consistent business-report style."""
from __future__ import annotations

import importlib.util
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path("docs/SafeAgent_Project_Guide.docx")
TABLE_HELPER = Path(r"C:\Users\User\.codex\plugins\cache\openai-primary-runtime\documents\26.715.12143\skills\documents\scripts\table_geometry.py")

spec = importlib.util.spec_from_file_location("table_geometry", TABLE_HELPER)
if spec is None or spec.loader is None:
    raise RuntimeError("Could not load document table geometry helper.")
table_geometry = importlib.util.module_from_spec(spec)
spec.loader.exec_module(table_geometry)

NAVY = "0B2545"
BLUE = "2E74B5"
BLUE_DARK = "1F4D78"
INK = "172B4D"
MUTED = "5E7189"
PALE = "F4F6F9"
PALE_BLUE = "E8EEF5"
GOLD = "7A5A00"
RED = "9B1C1C"
GREEN = "1F5D50"


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "D6DFEA") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def shade_paragraph(paragraph, fill: str, border_color: str = "CAD7E6") -> None:
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "6")
        element.set(qn("w:color"), border_color)
        borders.append(element)
    properties.append(borders)


def set_font(run, size: float, color: str = INK, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run("SafeAgent Project Guide | Page ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def paragraph(doc: Document, value: str = "", style: str | None = None, bold_prefix: str | None = None) -> object:
    item = doc.add_paragraph(style=style)
    if bold_prefix and value.startswith(bold_prefix):
        first = item.add_run(bold_prefix)
        set_font(first, 11, INK, bold=True)
        rest = item.add_run(value[len(bold_prefix):])
        set_font(rest, 11)
    else:
        run = item.add_run(value)
        set_font(run, 11)
    return item


def bullet(doc: Document, value: str) -> None:
    item = doc.add_paragraph(style="List Bullet")
    run = item.add_run(value)
    set_font(run, 11)


def numbered(doc: Document, value: str) -> None:
    item = doc.add_paragraph(style="List Number")
    run = item.add_run(value)
    set_font(run, 11)


def heading(doc: Document, value: str, level: int = 1) -> None:
    item = doc.add_paragraph(style=f"Heading {level}")
    run = item.add_run(value)
    set_font(run, {1: 16, 2: 13, 3: 12}[level], {1: BLUE, 2: BLUE, 3: BLUE_DARK}[level], bold=True)


def callout(doc: Document, label: str, body: str, color: str = PALE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    shade_paragraph(p, color)
    lead = p.add_run(label + " ")
    set_font(lead, 10.5, NAVY, bold=True)
    rest = p.add_run(body)
    set_font(rest, 10.5, INK)


def data_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header_cells = table.rows[0].cells
    set_repeat_table_header(table.rows[0])
    for cell, value in zip(header_cells, headers):
        set_cell_shading(cell, PALE_BLUE)
        set_cell_border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = cell.paragraphs[0].add_run(value)
        set_font(run, 10, NAVY, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            run = cell.paragraphs[0].add_run(value)
            set_font(run, 9.5, INK)
    table_geometry.apply_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_code(doc: Document, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    shade_paragraph(p, "EFF3F8", "D4DDE8")
    run = p.add_run(value)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(NAVY)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [("Heading 1", 16, BLUE, 18, 10), ("Heading 2", 13, BLUE, 12, 6), ("Heading 3", 12, BLUE_DARK, 8, 4)]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    footer = section.footer.paragraphs[0]
    for run in footer.runs:
        run.clear()
    add_page_number(footer)
    for run in footer.runs:
        set_font(run, 8.5, MUTED)
    return doc


def cover(doc: Document) -> None:
    doc.add_paragraph().paragraph_format.space_after = Pt(38)
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(8)
    run = kicker.add_run("AI SAFETY RESEARCH PROTOTYPE")
    set_font(run, 10, BLUE, bold=True)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("SafeAgent")
    set_font(run, 34, NAVY, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("A Human-Centered Command Safety Gateway for AI Coding Agents")
    set_font(run, 18, BLUE_DARK)
    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(20)
    run = intro.add_run("Technical and Conceptual Project Guide for Demonstrations, Research Discussions, and Graduate Scholarship Interviews")
    set_font(run, 12, MUTED, italic=True)
    callout(doc, "Purpose:", "This guide explains what SafeAgent does, how each component works, what evidence supports its claims, and the limits that must be stated honestly.", PALE_BLUE)
    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(26)
    meta.add_run("Prepared for project study and presentation\n").bold = True
    meta.add_run("Version: Hackathon research prototype | Date: July 2026\n")
    meta.add_run("Scope: Codex MCP command gateway, local human approval, audit evidence, and research-quality reporting")
    for run in meta.runs:
        set_font(run, 10.5, MUTED)
    doc.add_page_break()


def build() -> None:
    doc = setup_document()
    cover(doc)

    heading(doc, "How to Use This Guide")
    paragraph(doc, "Read the first three sections if you need a quick explanation for a non-technical audience. Read Sections 4 through 10 if you need to defend the architecture, research choices, testing, and limitations in front of technical reviewers.")
    data_table(doc, ["Audience", "Recommended sections", "What they should understand"], [
        ["Non-technical reviewer", "1-3, 8", "The problem, the human approval flow, and the project value."],
        ["Professor or researcher", "4-7, 9-10", "Risk modeling, held-out certification, evidence, and limitations."],
        ["Hackathon judge", "2, 5, 8, 11", "The live demo, architecture, safety controls, and honest claim."],
        ["Project owner", "All sections", "How to run, explain, improve, and defend the project."],
    ], [2000, 2500, 4860])

    heading(doc, "1. The Problem SafeAgent Solves")
    paragraph(doc, "AI coding agents can inspect repositories, write files, and propose terminal commands. That makes them useful, but it also creates a safety problem: a command can be syntactically valid and still delete files, overwrite disks, change permissions, remove packages, or execute unreviewed code downloaded from the internet.")
    paragraph(doc, "The core question is not whether a user is allowed to run a dangerous command. The user remains in control. The question is whether an AI-generated command reaches the operating system without a clear risk explanation and an explicit human decision.")
    callout(doc, "SafeAgent in one sentence:", "It is a command execution gateway that evaluates an AI agent's requested shell command before execution, explains the risk, obtains human approval when needed, and records what happened.")
    heading(doc, "Why this matters", level=2)
    bullet(doc, "A coding agent can act quickly across many files; one mistaken command can have a large impact.")
    bullet(doc, "Terminal syntax is compact, so destructive behavior can be hard to notice in a long command.")
    bullet(doc, "Human approval is most useful when it arrives with a clear explanation, not only a warning score.")
    bullet(doc, "Researchers need evidence about a gate's behavior rather than an unsupported claim that it is universally safe.")

    heading(doc, "2. The SafeAgent Idea in Plain Language")
    paragraph(doc, "Think of SafeAgent as a security checkpoint between an AI agent and the computer. The agent does not directly ask the operating system to run a command through the SafeAgent path. Instead, it asks the SafeAgent MCP tool to execute the command. SafeAgent checks the command, decides whether a human must review it, and only then calls the shell.")
    numbered(doc, "An AI coding agent proposes a command, such as git status or rm -rf build.")
    numbered(doc, "The SafeAgent MCP server receives the command through execute_command(command).")
    numbered(doc, "The deterministic rules and optional machine-learning model calculate risk features and a final score.")
    numbered(doc, "Safe commands execute immediately. Risky and dangerous commands are sent to the local approval console.")
    numbered(doc, "The user sees the command, score, category, reasons, and optional AI advisory, then chooses Yes or No.")
    numbered(doc, "SafeAgent executes only an allowed command and writes an audit event for every decision.")
    add_code(doc, "AI agent -> SafeAgent MCP -> risk assessment -> human approval (when needed) -> OS execution -> JSONL audit log -> dashboard")

    heading(doc, "3. Architecture and Component Responsibilities")
    data_table(doc, ["Component", "Responsibility", "Why it exists"], [
        ["safeagent/mcp_server.py", "Exposes execute_command to MCP clients and returns structured execution results.", "Makes SafeAgent available as an execution tool to Codex and other compatible agents."],
        ["safeagent/gate.py", "Coordinates scoring, approval, platform-aware execution, timing, and fail-closed handling.", "Creates one auditable decision boundary."],
        ["safeagent/cost_model.py", "Matches inspectable destructive-command patterns.", "Provides fast, deterministic, explainable protection."],
        ["safeagent/risk_models.py", "Combines rule score with optional logistic ML probability.", "Allows supervised risk signals without removing rule authority."],
        ["safeagent/approval_console.py", "Runs the separate human confirmation UI over localhost.", "Keeps interactive input out of MCP stdio streams."],
        ["safeagent/logging_store.py", "Writes append-only JSON Lines audit events.", "Supports transparency, debugging, and dashboard analysis."],
        ["safeagent/certification.py", "Computes a held-out Hoeffding upper confidence bound.", "Separates safety evidence from intuition or marketing claims."],
        ["dashboard.py", "Displays operational and research evidence from the log and model metadata.", "Makes safety activity understandable to users and reviewers."],
    ], [2400, 3600, 3360])
    callout(doc, "Important design rule:", "The GPT advisory can improve explanation quality, but it never chooses allow or block. The safety decision comes from the risk model and the human approval result.", "FFF5DB")

    heading(doc, "4. Command Risk Assessment")
    heading(doc, "4.1 Deterministic rule-based model", level=2)
    paragraph(doc, "The rule-based model is the always-available foundation. It uses conservative regular-expression patterns that are easy to inspect and edit. A matched rule produces a score, a rule identifier, and a plain-language explanation.")
    data_table(doc, ["Example behavior", "Typical category", "Why review is needed"], [
        ["ls, pwd, git status, python app.py", "Safe", "These commands are usually read-only or low-impact in normal project use."],
        ["git push --force, pip uninstall, mv important_file", "Risky", "They can overwrite collaboration history, remove dependencies, or move important data."],
        ["rm -rf, rmdir /s, Remove-Item -Recurse, find -delete", "Dangerous", "They can permanently remove many files."],
        ["mkfs, dd ... of=/dev, chmod -R 777", "Dangerous", "They can destroy storage or weaken security boundaries."],
        ["curl or wget piped to bash or sh", "Dangerous", "They download code and execute it without inspection."],
    ], [3200, 1600, 4560])
    heading(doc, "4.2 Optional logistic machine-learning model", level=2)
    paragraph(doc, "SafeAgent also contains an extensible supervised risk-model interface. The included implementation is dependency-free logistic regression trained on labeled shell commands. It predicts the probability that a command belongs to the unsafe class. The model is optional: if no trained model artifact is present, the system remains rule-based rather than inventing a probability.")
    paragraph(doc, "Feature engineering is transparent. The model observes command length, token count, shell operators, recursive actions, deletion, overwriting, privilege escalation, network access, download-and-execute patterns, encoding techniques, suspicious paths, process termination, and package-management behavior.")
    heading(doc, "4.3 Hybrid final score", level=2)
    paragraph(doc, "When a trained ML model is loaded, SafeAgent combines the deterministic score r and ML unsafe probability m with a conservative noisy-OR rule: final_score = 1 - (1 - r)(1 - m). This means ML can raise a score, but it cannot reduce a score already assigned by a deterministic safety rule.")
    data_table(doc, ["Final score", "Category", "Execution behavior"], [
        ["Below 0.40", "SAFE", "Automatic execution through the SafeAgent path."],
        ["0.40 to below 0.85", "RISKY", "Human approval is required."],
        ["0.85 or higher", "DANGEROUS", "Human approval is required with detailed explanation."],
    ], [2100, 2100, 5160])

    heading(doc, "5. Human Approval, Execution, and Fail-Closed Behavior")
    paragraph(doc, "For risky and dangerous commands, the MCP server contacts a local approval console. The console displays the exact command, category, risk score, deterministic reasons, and any available AI-generated advisory. The user then chooses Y to execute or N to cancel.")
    paragraph(doc, "The approval console is separate from the MCP process because standard MCP over stdio reserves the server's input and output for protocol messages. If the MCP server tried to call input() directly, it could corrupt the protocol. The localhost approval bridge is therefore an architectural necessity, not a cosmetic UI choice.")
    data_table(doc, ["Failure condition", "SafeAgent response", "Why this is safer"], [
        ["Parser, cost model, logging preflight, MCP routing, or configured advisory fails", "BLOCKED_ERROR; no execution", "A missing safety check never becomes permission to run a command."],
        ["Approval console unavailable", "BLOCKED_ERROR; no execution", "A human decision is required but cannot be obtained."],
        ["User presses N", "CANCELLED; no execution", "The user retains final control."],
        ["User does not respond", "Approval times out after configurable default of 300 seconds; no execution", "Avoids a stuck request while remaining fail-closed."],
        ["User presses Y", "Execution begins immediately", "There is no arbitrary 120-second execution deadline after approval."],
    ], [3300, 2750, 3310])
    paragraph(doc, "For Windows usability, SafeAgent only translates harmless aliases such as pwd to cd and ls to dir. It never silently rewrites risky or dangerous commands.")

    heading(doc, "6. MCP Integration: What It Proves and What It Does Not")
    paragraph(doc, "SafeAgent uses an MCP stdio server with one primary tool: execute_command(command: str). The project contains a real integration test that starts the MCP server, completes the protocol handshake, lists the tool schema, calls a safe command, and confirms that an audit log was written.")
    add_code(doc, "python -m pytest tests/test_mcp_stdio_integration.py -q")
    callout(doc, "Critical enforcement boundary:", "MCP is additive. When Codex calls execute_command, SafeAgent controls that command. However, a stock Codex CLI may still have its own native shell tool. MCP registration alone cannot disable that built-in executor. Therefore SafeAgent must not claim universal interception of stock Codex shell actions.", "FDEFF2")
    paragraph(doc, "A truly non-bypassable deployment would require either a supported Codex native-shell hook or disable setting, or an agent runtime designed with SafeAgent as its only execution tool. Until then, SafeAgent should be used alongside Codex's native sandbox and approval controls.")
    heading(doc, "Reliable registration procedure", level=2)
    numbered(doc, "Create and activate the project virtual environment, then install the project in editable mode.")
    numbered(doc, "Start the approval console in a separate terminal.")
    numbered(doc, "Register MCP using an absolute path to the virtual-environment Python executable.")
    numbered(doc, "End the current Codex session and open a new session, because MCP tools are discovered at session startup.")
    add_code(doc, "$python = (Resolve-Path .\\.venv\\Scripts\\python.exe).Path\ncodex mcp remove safeagent\ncodex mcp add safeagent -- $python -m safeagent.mcp_server\ncodex mcp list")

    heading(doc, "7. Audit Logging and Dashboard Evidence")
    paragraph(doc, "Every gate decision is stored as a JSON Lines event. New-format records include timestamp, operating system, original and executed command, category, rule score, ML probability, final score, matched rules, feature values, decision, user response, approval time, execution time, standard output, standard error, advisory text, and errors.")
    paragraph(doc, "The dashboard reads those records defensively. Malformed events are ignored for visualizations instead of crashing the page. The simplified Security Operations Center view provides executive KPIs, service health, a risk timeline, a score distribution, decision distribution, readable recent activity, and a focused audit explorer. CSV and JSON exports preserve the filtered audit view.")
    callout(doc, "Interpretation note:", "The dashboard shows operational evidence. It does not prove that all commands are safe, and it does not turn a visual chart into a statistical guarantee.")

    heading(doc, "8. HC-RLHF Inspiration and Statistical Certification")
    paragraph(doc, "SafeAgent is inspired by selected ideas from High-Confidence Safe Reinforcement Learning from Human Feedback (HC-RLHF). The paper separates helpfulness from harmlessness, uses a pessimistic safety constraint during optimization, and performs a separate safety test using an upper confidence bound. SafeAgent adapts the safety-side ideas to shell commands.")
    data_table(doc, ["HC-RLHF paper concept", "SafeAgent adaptation", "Implemented?"], [
        ["Language-model action", "A shell command requested by an AI coding agent", "Yes, as an application mapping."],
        ["Cost model", "Rule-based score plus optional logistic unsafe probability", "Yes, as a command-risk model."],
        ["Human feedback", "Approve or cancel a risky command", "Yes, as a real-time execution decision."],
        ["Held-out safety test", "Hoeffding upper bound on unsafe auto-approvals", "Yes, as a certification framework."],
        ["RLHF reward model", "No equivalent helpfulness model", "No."],
        ["Full Seldonian optimization", "No candidate-selection/optimization loop", "No."],
    ], [2800, 4200, 2360])
    paragraph(doc, "The held-out test measures unsafe commands that the gate would automatically approve. If p_hat is the observed unsafe-auto-approval rate, N is the number of calibration samples, and delta is the allowed failure probability, SafeAgent reports the one-sided Hoeffding bound: p_hat + sqrt(log(1/delta) / (2N)).")
    callout(doc, "Current honest result:", "With the included calibration data, N = 24, unsafe auto-approvals = 0, observed rate = 0.00%, and the 95% upper bound is 24.98%. With a 5.00% target, the system is NOT CERTIFIED. The sample is too small to support a strong claim.", "FFF5DB")
    paragraph(doc, "When a future calibration run does meet its threshold, the appropriate claim is: with the stated confidence, unsafe auto-approvals remain below the configured threshold on the held-out calibration distribution. It is not a universal guarantee over every command, operating system, or attacker strategy.")

    heading(doc, "9. Training, Evaluation, and Reproducibility")
    paragraph(doc, "The project keeps demonstration training data separate from calibration data. Training data fits the optional logistic model. Calibration data is reserved for the held-out safety test. This separation avoids using the same examples both to tune a system and to claim that it passed an independent test.")
    add_code(doc, "python train_model.py --dataset data/training.csv --model-out models/logistic_risk_model.json\npython evaluate_model.py --dataset data/training.csv --model models/logistic_risk_model.json\npython calibrate_model.py --calibration data/calibration.csv --model models/logistic_risk_model.json")
    paragraph(doc, "The evaluation script reports accuracy, precision, recall, F1 score, false positives, and false negatives. These metrics describe performance on the labeled dataset; they do not replace the held-out unsafe-auto-approval certification test.")
    heading(doc, "How to explain the difference", level=2)
    bullet(doc, "Accuracy asks: how often did the classifier match the labels?")
    bullet(doc, "Precision asks: when it predicts unsafe, how often is it correct?")
    bullet(doc, "Recall asks: how many unsafe examples did it catch?")
    bullet(doc, "False negatives are especially important because they are unsafe commands missed by the model.")
    bullet(doc, "Certification asks a different question: how large could the unsafe auto-approval rate be, with high confidence, on held-out data?")

    heading(doc, "10. Demo Script for Judges, Researchers, and Professors")
    heading(doc, "Before the demo", level=2)
    numbered(doc, "Run the test suite: python -m pytest -q.")
    numbered(doc, "Optionally train the logistic model, then start the Streamlit dashboard.")
    numbered(doc, "Start the approval console: python -m safeagent.approval_console.")
    numbered(doc, "Register the MCP server with the absolute-path procedure and begin a fresh Codex session.")
    heading(doc, "Live story", level=2)
    numbered(doc, "Ask the agent to use SafeAgent execute_command for git status. Explain that this is a low-risk read operation and it executes automatically.")
    numbered(doc, "Ask the agent to call execute_command for rmdir /s disposable_demo_folder or rm -rf disposable_demo_folder. Use only a folder created for the demonstration.")
    numbered(doc, "Show the SafeAgent approval console: category, risk score, deterministic reason, and advisory if configured.")
    numbered(doc, "Enter N. Show that the folder remains and that the dashboard records a cancellation.")
    numbered(doc, "Repeat only with a disposable test folder and enter Y. Show the approval, execution result, timing, and audit event.")
    numbered(doc, "Open the certification panel and explain why NOT CERTIFIED is an honest, positive research practice rather than a product failure.")
    callout(doc, "Do not overclaim:", "Say that the MCP tool controls commands routed through it. Also say that stock Codex native shell execution remains a separate bypass boundary until a supported native hook or restricted runtime is available.", "FDEFF2")

    heading(doc, "11. Testing and Verification Evidence")
    data_table(doc, ["Area", "Verification approach", "What it demonstrates"], [
        ["Rules and parser", "Unit tests for safe, risky, dangerous, empty, and malformed commands", "Risk categories and fail-closed parser behavior."],
        ["Hybrid model", "Feature, logistic-model persistence, and score-conservatism tests", "ML signals can raise but not lower deterministic risk."],
        ["MCP service", "Service-level tests plus a real stdio handshake test", "Tool registration, schema discovery, safe execution path, and audit writing."],
        ["Approval console", "Socket client, response, and health-probe tests", "Human confirmation transport works without corrupting MCP stdio."],
        ["Execution layer", "Platform-normalization tests and mock executor tests", "Safe aliases translate; dangerous commands are not silently transformed."],
        ["Dashboard", "Headless startup and malformed-audit rendering checks", "Malformed events do not crash user-facing visualizations."],
        ["Current suite", "Project virtual environment test run", "31 automated tests passed at the latest verification."],
    ], [2150, 3900, 3310])

    heading(doc, "12. Limitations, Threat Model, and Roadmap")
    heading(doc, "Current limitations", level=2)
    bullet(doc, "Regex rules and a small demonstration dataset do not form a complete shell parser or malware detector.")
    bullet(doc, "The logistic model is optional and should not be treated as a production classifier without larger, independently reviewed datasets.")
    bullet(doc, "The 24-sample calibration set is insufficient for a 5% unsafe-auto-approval certification target.")
    bullet(doc, "MCP controls only commands sent through SafeAgent; it does not automatically disable a stock Codex native shell tool.")
    bullet(doc, "The local approval console is designed for one trusted user on localhost, not for multi-user enterprise authorization.")
    heading(doc, "Future work", level=2)
    bullet(doc, "Collect larger, independently curated command datasets and preserve a truly held-out calibration split.")
    bullet(doc, "Evaluate richer risk models such as gradient-boosted trees or neural encoders through the existing RiskModel interface.")
    bullet(doc, "Add robust shell parsing, command decomposition, filesystem context, and least-privilege execution sandboxes.")
    bullet(doc, "Integrate with a supported native Codex shell hook or a controlled agent runtime to remove the direct-shell bypass boundary.")
    bullet(doc, "Build a full research pipeline only if it includes proper governance, offline evaluation, and new held-out safety tests; do not add online adaptation casually.")

    heading(doc, "13. Questions You May Be Asked")
    data_table(doc, ["Question", "Strong concise answer"], [
        ["Is this an AI model that blocks commands?", "It is a safety gate. It combines interpretable rules with an optional learned risk estimate, then keeps a human in control for risky commands."],
        ["Does GPT decide whether to execute?", "No. GPT is advisory only. It can explain danger and propose a safer alternative, but the gate and human approval decide execution."],
        ["Did you implement HC-RLHF?", "No. The project is inspired by selected HC-RLHF safety principles: decoupled cost modeling, conservative safety constraints, and held-out confidence-bound testing."],
        ["Is the system certified safe?", "Not with the included small calibration set. The dashboard honestly reports NOT CERTIFIED because the 95% upper bound is above the 5% target."],
        ["Can Codex bypass it?", "A command routed through the MCP tool cannot bypass SafeAgent. Stock Codex's separate native shell remains a documented limitation until a native hook or controlled runtime is used."],
        ["Why not simply block dangerous commands?", "The design preserves human control. It explains the consequences and asks the human to decide, rather than silently taking authority away."],
    ], [2900, 6460])

    heading(doc, "14. Glossary")
    data_table(doc, ["Term", "Meaning in SafeAgent"], [
        ["AI coding agent", "Software such as Codex that can reason about a codebase and request tool actions."],
        ["MCP", "Model Context Protocol, the tool-connection mechanism used to expose execute_command."],
        ["Risk score", "A number from 0 to 1 used to categorize a command as safe, risky, or dangerous."],
        ["Cost model", "A model of harmfulness or risk; in SafeAgent it estimates command danger, not usefulness."],
        ["Fail closed", "If a safety component fails, the command is not executed."],
        ["Calibration set", "Held-out labeled examples used to estimate a safety bound rather than train the model."],
        ["Hoeffding bound", "A statistical upper-confidence bound used here to state a conservative unsafe-auto-approval estimate."],
        ["False negative", "An unsafe command that a classifier incorrectly treats as safe; a key safety concern."],
    ], [2500, 6860])

    heading(doc, "15. References")
    paragraph(doc, "Chittepu, Y., Metevier, B., Schwarzer, W., Hoag, A., Niekum, S., and Thomas, P. S. (2025). Reinforcement Learning from Human Feedback with High-Confidence Safety Constraints. Reinforcement Learning Conference. https://arxiv.org/abs/2506.08266")
    paragraph(doc, "SafeAgent source modules, tests, dashboard, training data, calibration data, and README in the local project repository.")
    paragraph(doc, "End of guide.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
